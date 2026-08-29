"""文档解析。

支持文本型 PDF / Markdown / TXT / XLSX / CSV / DOCX 的文本抽取。
重型依赖（PyMuPDF / python-docx）按需 lazy import，缺失时抛出 ``ParseError`` 并提示安装方式；
md / txt / csv 走标准库，零依赖，保证开发与测试环境可用性。
"""

from __future__ import annotations

import io
import os
import threading
import zipfile
from dataclasses import dataclass

import config


class ParseError(Exception):
    """文档解析失败。"""


@dataclass
class ParsedDocument:
    text: str
    source: str  # 使用的解析器名称，便于排查


_OFFICE_EXTENSIONS = {".xlsx", ".xlsm", ".docx"}
_GENERIC_MIME_TYPES = {"", "application/octet-stream", "binary/octet-stream"}
_MIME_TYPES = {
    ".md": {"text/markdown", "text/plain"},
    ".markdown": {"text/markdown", "text/plain"},
    ".txt": {"text/plain"},
    ".csv": {"text/csv", "application/csv", "application/vnd.ms-excel"},
    ".pdf": {"application/pdf"},
    ".xlsx": {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    },
    ".xlsm": {"application/vnd.ms-excel.sheet.macroenabled.12"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
}
_DOCUMENT_PARSE_SEM = threading.BoundedSemaphore(config.UPLOAD_GLOBAL_CONCURRENCY)


def _validate_office_zip(raw: bytes, ext: str) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            infos = archive.infolist()
            if len(infos) > config.UPLOAD_MAX_ZIP_ENTRIES:
                raise ParseError(
                    f"Office 文件 ZIP 条目超过 {config.UPLOAD_MAX_ZIP_ENTRIES} 个上限"
                )
            total_size = sum(info.file_size for info in infos)
            if total_size > config.UPLOAD_MAX_DECOMPRESSED_BYTES:
                limit_mb = config.UPLOAD_MAX_DECOMPRESSED_BYTES // (1024 * 1024)
                raise ParseError(f"Office 文件解压后超过 {limit_mb}MB 上限")
            names = {info.filename.replace("\\", "/") for info in infos}
            if any(
                name.startswith("/") or ".." in name.split("/")
                for name in names
            ):
                raise ParseError("Office 文件包含不安全的 ZIP 路径")
            if any(info.flag_bits & 0x1 for info in infos):
                raise ParseError("不支持加密的 Office 文件")
            required = "word/document.xml" if ext == ".docx" else "xl/workbook.xml"
            if "[Content_Types].xml" not in names or required not in names:
                raise ParseError(f"文件内容不是有效的 {ext[1:].upper()} 文档")
    except ParseError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise ParseError(f"文件内容不是有效的 {ext[1:].upper()} 文档") from exc


def validate_upload(filename: str, raw: bytes, content_type: str | None = None) -> None:
    """校验上传文件的扩展名、声明 MIME 与基础文件签名。"""
    ext = os.path.splitext(filename or "")[1].lower()
    if not ext:
        raise ParseError("文件必须包含扩展名")
    if ext not in _MIME_TYPES:
        raise ParseError(f"不支持的文件类型：{ext}")
    mime = (content_type or "").split(";", 1)[0].strip().lower()
    if mime not in _GENERIC_MIME_TYPES and mime not in _MIME_TYPES[ext]:
        raise ParseError(f"文件 MIME 类型 {mime or '未知'} 与扩展名 {ext} 不匹配")
    if ext == ".pdf":
        if not raw.startswith(b"%PDF-"):
            raise ParseError("文件内容不是有效的 PDF")
    elif ext in _OFFICE_EXTENSIONS:
        if not raw.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
            raise ParseError(f"文件内容不是有效的 {ext[1:].upper()} 文档")
        _validate_office_zip(raw, ext)
    else:
        sample = raw[:8192]
        if sample.startswith((b"%PDF-", b"PK\x03\x04")) or b"\x00" in sample:
            raise ParseError(f"文件内容与文本扩展名 {ext} 不匹配")


def _parse_plain(raw: bytes) -> str:
    # md / txt：直接读文本
    return raw.decode("utf-8", errors="replace")


def _parse_csv(raw: bytes) -> str:
    text = raw.decode("utf-8", errors="replace")
    lines = [ln.rstrip() for ln in text.splitlines() if ln.strip()]
    return "\n".join(lines)


def _parse_pdf(raw: bytes | None, path: str | None) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover - 依赖可选
        raise ParseError("解析 PDF 需要 PyMuPDF：pip install pymupdf") from exc
    doc = fitz.open(stream=raw, filetype="pdf") if raw is not None else fitz.open(path)
    try:
        if doc.page_count > config.UPLOAD_MAX_PDF_PAGES:
            raise ParseError(f"PDF 超过 {config.UPLOAD_MAX_PDF_PAGES} 页上限")
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _parse_xlsx(raw: bytes | None, path: str | None) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ParseError("解析 XLSX 需要 openpyxl（已包含在 requirements）") from exc
    src: io.BytesIO | str = io.BytesIO(raw) if raw is not None else path  # type: ignore[assignment]
    wb = load_workbook(filename=src, data_only=True, read_only=True)
    chunks: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            line = "\t".join(cells).strip()
            if line:
                chunks.append(line)
    return "\n".join(chunks)


def _parse_docx(raw: bytes | None, path: str | None) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ParseError("解析 DOCX 需要 python-docx：pip install python-docx") from exc
    src: io.BytesIO | str = io.BytesIO(raw) if raw is not None else path  # type: ignore[assignment]
    document = Document(src)
    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


def _parse_document(
    filename: str,
    raw: bytes | None = None,
    path: str | None = None,
) -> ParsedDocument:
    """按扩展名选择解析器。优先使用 ``raw`` 字节，否则读取 ``path``。"""
    if raw is None and path is None:
        raise ParseError("parse_document 需要 raw 字节或文件路径")
    src_path = path or filename
    ext = os.path.splitext(src_path)[1].lower()
    try:
        if ext in (".md", ".markdown", ".txt"):
            parsed = ParsedDocument(_parse_plain(raw or b""), "plain")
        elif ext == ".csv":
            parsed = ParsedDocument(_parse_csv(raw or b""), "csv")
        elif ext == ".pdf":
            parsed = ParsedDocument(_parse_pdf(raw, path), "pymupdf")
        elif ext in (".xlsx", ".xlsm"):
            if raw is not None:
                _validate_office_zip(raw, ext)
            parsed = ParsedDocument(_parse_xlsx(raw, path), "openpyxl")
        elif ext == ".docx":
            if raw is not None:
                _validate_office_zip(raw, ext)
            parsed = ParsedDocument(_parse_docx(raw, path), "python-docx")
        else:
            raise ParseError(f"不支持的文件类型：{ext or '未知'}")
        if len(parsed.text) > config.UPLOAD_MAX_EXTRACTED_CHARS:
            raise ParseError(
                f"文档抽取文本超过 {config.UPLOAD_MAX_EXTRACTED_CHARS} 字符上限"
            )
        return parsed
    except ParseError:
        raise
    except Exception as exc:  # pragma: no cover - 解析内部异常统一包装
        raise ParseError(f"解析 {src_path} 失败：{exc}") from exc


def parse_document(
    filename: str,
    raw: bytes | None = None,
    path: str | None = None,
) -> ParsedDocument:
    """串行化重型文档解析；上传、重索引和知识沉淀共用同一进程级额度。"""
    with _DOCUMENT_PARSE_SEM:
        return _parse_document(filename, raw=raw, path=path)
